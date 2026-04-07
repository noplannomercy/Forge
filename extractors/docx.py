import io
from docx import Document

from models import ConvertResult, Quality


def _table_to_md(table) -> str:
    """docx 표 → markdown 표 변환"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 1:
        header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
        rows.insert(1, header_sep)
    return "\n".join(rows)


async def extract(file_bytes: bytes, file_name: str) -> ConvertResult:
    """DOCX → Markdown 변환"""
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # namespace 제거

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name:
                parts.append(f"# {text}")
            elif "Heading 2" in style_name:
                parts.append(f"## {text}")
            elif "Heading 3" in style_name:
                parts.append(f"### {text}")
            else:
                parts.append(text)

        elif tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    parts.append(_table_to_md(table))
                    break

    full_text = "\n\n".join(parts)
    total_chars = len(full_text.strip())

    return ConvertResult(
        text=full_text,
        format="md",
        pages=1,
        file_name=file_name,
        source_format="docx",
        route="extract",
        quality=Quality(
            total_chars=total_chars,
            chars_per_page=total_chars if total_chars > 0 else 0,
            total_pages=1,
            failed_pages=0,
            confidence="high",
        ),
    )
